"""资料解析：支持 txt / pdf / docx / 图片(OCR) 四种格式的文本提取。

OCR 后端（可选，Apache 2.0, https://github.com/PaddlePaddle/PaddleOCR）：
- 扫描版 PDF / 图片真题（jpg/png）自动走 PaddleOCR 识别
- 文本版 PDF 仍用 pdfplumber（MIT）快速提取
- 设置 USE_OCR=1 启用；未安装 paddleocr 时给出提示
"""
import os


def parse_file(path: str) -> str:
    """按扩展名解析文件为纯文本。"""
    if not os.path.isfile(path):
        raise FileNotFoundError(f"文件不存在: {path}")
    lower = path.lower()
    if lower.endswith(".txt"):
        return _parse_txt(path)
    if lower.endswith(".pdf"):
        return _parse_pdf(path)
    if lower.endswith(".docx"):
        return _parse_docx(path)
    if lower.endswith((".jpg", ".jpeg", ".png", ".bmp", ".webp")):
        return _parse_image_ocr(path)
    raise ValueError(f"不支持的文件格式: {path}（支持 txt/pdf/docx/图片OCR）")


def _parse_txt(path: str) -> str:
    # utf-8-sig 优先：能正确处理带 BOM 的 UTF-8 文件
    for enc in ("utf-8-sig", "utf-8", "gbk"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文件编码: {path}")


def _parse_pdf(path: str) -> str:
    """PDF 解析：先尝试文本层（pdfplumber），无文本层时降级 PaddleOCR 逐页识别。"""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("解析 PDF 需要安装 pdfplumber: pip install pdfplumber")
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    result = "\n".join(parts)
    if not result.strip():
        # 扫描版 PDF → OCR 兜底
        ocr_text = _pdf_ocr_fallback(path)
        if ocr_text.strip():
            return ocr_text
        raise ValueError(f"PDF 文件无可提取文本（可能是扫描件）: {path}")
    return result


def _pdf_ocr_fallback(path: str) -> str:
    """扫描版 PDF：用 PaddleOCR 逐页识别（需安装 paddleocr + pymupdf 渲染）。"""
    if os.getenv("USE_OCR") != "1":
        return ""
    try:
        from paddleocr import PaddleOCR
        import fitz  # PyMuPDF 渲染页面为图片（AGPL 双许可，仅运行期使用）

        ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        parts = []
        with fitz.open(path) as doc:
            for page in doc:
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                import io
                result = ocr.ocr(io.BytesIO(img_bytes), cls=True)
                for line in result or []:
                    for item in line or []:
                        parts.append(item[1][0])
        return "\n".join(parts)
    except ImportError:
        return ""


def _parse_docx(path: str) -> str:
    try:
        import docx
    except ImportError:
        raise RuntimeError("解析 docx 需要安装 python-docx: pip install python-docx")
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    if not parts:
        # 尝试从表格中提取
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    if not parts:
        raise ValueError(f"docx 文件无可提取文本: {path}")
    return "\n".join(parts)


def _parse_image_ocr(path: str) -> str:
    """图片真题 OCR 识别（PaddleOCR，Apache 2.0 许可证）。"""
    if os.getenv("USE_OCR") != "1":
        raise RuntimeError(
            "图片解析需要 OCR：设置 USE_OCR=1 并安装 paddleocr "
            "(pip install paddleocr)，或先把图片转成 txt"
        )
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        result = ocr.ocr(path, cls=True)
        lines = []
        for line in result or []:
            for item in line or []:
                lines.append(item[1][0])
        text = "\n".join(lines)
        if not text.strip():
            raise ValueError(f"OCR 未识别到文本: {path}")
        return text
    except ImportError as e:
        raise RuntimeError(f"PaddleOCR 不可用: {e}。安装: pip install paddleocr")
